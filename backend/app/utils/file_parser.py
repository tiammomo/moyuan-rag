"""
文件解析工具 - 将不同格式文件转换为文本
基于 src/file_utils.py 的 FileParser 类
"""
import logging
import traceback
import re
import os
import time
import gc
from pathlib import Path
from typing import BinaryIO, Optional, List, Dict, Any
from concurrent.futures import ProcessPoolExecutor, as_completed

import fitz as pymupdf  # PyMuPDF < 1.24.0 使用 fitz 导入
import docx
import html2text
from bs4 import BeautifulSoup

logger = logging.getLogger(__name__)

# --- 性能优化常量 ---
MAX_PAGE_BLOCKS = 3000      # 单页最大块数量，超过则认为是复杂矢量图或异常数据，进行截断或简化处理
MAX_PAGE_CHARS = 50000      # 单页最大字符数量限制
SINGLE_PAGE_TIMEOUT = 10.0  # 单页解析硬超时（秒）
MEMORY_THRESHOLD = 512      # 内存阈值 (MB)，超过则强制执行 GC
# --------------------

def _parse_pdf_pages_worker(file_path: str, start_page: int, end_page: int, h2t_config: dict) -> str:
    """子进程任务：解析 PDF 的指定页面范围 (性能增强版)"""
    pid = os.getpid()
    start_t = time.time()
    logger.info(f"[Process {pid}] 开始解析 PDF 任务范围: {start_page} - {end_page}")
    
    try:
        # 使用 garbage=4 减少内存占用，适用于大文件
        doc = pymupdf.open(file_path)
        md_content = []
        
        # 预先初始化 html2text 仅作为备用
        h2t = html2text.HTML2Text()
        for key, value in h2t_config.items():
            setattr(h2t, key, value)
            
        for i in range(start_page, end_page):
            page_start_t = time.time()
            try:
                page = doc[i]
                
                # 1. 异常数据识别：检查页面是否有过多的矢量图形（这些通常导致解析极其缓慢）
                # PyMuPDF 的 get_drawings() 如果非常多，说明是复杂的矢量图
                # 但 get_drawings 比较慢，我们先通过 get_text("blocks") 的耗时来判断
                
                # 优先尝试 blocks 模式
                blocks = page.get_text("blocks")
                block_count = len(blocks)
                
                # 2. 异常数据处理：如果块数量异常多，可能是由于复杂的 CAD 图纸或损坏的文本层
                if block_count > MAX_PAGE_BLOCKS:
                    logger.warning(f"[Process {pid}] Page {i} 块数量过多 ({block_count}), 触发异常数据降级处理")
                    # 降级方案：直接提取纯文本，不保留块结构，避免排序和复杂处理
                    md_text = page.get_text("text")
                else:
                    # 正常处理：按照阅读顺序排序
                    blocks.sort(key=lambda b: (b[1], b[0]))
                    
                    page_text = []
                    total_chars = 0
                    for b in blocks:
                        text = b[4].strip()
                        if text:
                            # 过滤掉图片块
                            if b[6] == 1: 
                                continue
                            
                            # 3. 字符数限制：防止某些异常 PDF 产生无限重复文本
                            if total_chars + len(text) > MAX_PAGE_CHARS:
                                logger.warning(f"[Process {pid}] Page {i} 字符数超限，进行截断")
                                break
                                
                            page_text.append(text)
                            total_chars += len(text)
                    
                    md_text = "\n\n".join(page_text)
                    
                    # 如果 blocks 模式提取不到内容，再尝试 HTML 模式（兜底）
                    if not md_text.strip():
                        html_content = page.get_text("html")
                        md_text = h2t.handle(html_content)
                
            except Exception as e:
                logger.warning(f"[Process {pid}] Page {i} 解析失败，尝试最简文本提取: {e}")
                try:
                    md_text = doc[i].get_text("text")
                except:
                    md_text = ""

            md_content.append(FileParser._format_page_block(i + 1, md_text))
            
            # 4. 性能审计日志
            page_duration = time.time() - page_start_t
            if page_duration > 2.0:
                logger.warning(f"[Process {pid}] Page {i} 耗时: {page_duration:.2f}s | 块数: {len(blocks) if 'blocks' in locals() else 'N/A'} | 长度: {len(md_text)}")
            else:
                logger.debug(f"[Process {pid}] Page {i} 完成: {page_duration:.2f}s")
            
            # 5. 定期清理内存
            if i % 20 == 0:
                gc.collect()

        doc.close()
        duration = time.time() - start_t
        logger.info(f"[Process {pid}] 完成任务范围: {start_page} - {end_page} | 总耗时: {duration:.2f}s")
        return "\n\n".join(md_content)
    except Exception as e:
        error_stack = traceback.format_exc()
        logger.error(f"[Process {pid}] 致命错误 {start_page}-{end_page}:\n{error_stack}")
        raise


class FileParser:
    """单文件解析器：将不同格式文件转换为纯文本"""

    def __init__(self):
        # 配置 HTML 转 Markdown 的参数
        self.h2t = html2text.HTML2Text()
        self.h2t.ignore_links = False
        self.h2t.ignore_images = True
        self.h2t.body_width = 0
        self.h2t.ignore_emphasis = False

    @staticmethod
    def _format_page_block(page_number: int, content: str) -> str:
        marker = f"<<<PAGE:{page_number}>>>"
        content = (content or "").strip()
        if not content:
            return marker
        return f"{marker}\n{content}"

    def parse_file(self, file_path: Path) -> str:
        """根据文件后缀解析文件"""
        suffix = file_path.suffix.lower()
        try:
            if suffix == '.docx':
                return self._parse_docx(file_path)
            elif suffix == '.pdf':
                return self._parse_pdf(file_path)
            elif suffix == '.html':
                return self._parse_html(file_path)
            elif suffix == '.md':
                return self._parse_md(file_path)
            elif suffix == '.txt':
                return self._parse_txt(file_path)
            else:
                logger.warning(f"不支持的文件格式: {file_path}")
                return ""
        except Exception as e:
            error_detail = traceback.format_exc()
            logger.error(f"解析文件失败 {file_path}:\n{error_detail}")
            raise

    def _parse_docx(self, file_path: Path) -> str:
        """解析 .docx 文件，保留标题层级和表格"""
        doc = docx.Document(file_path)
        md_lines = []

        p_idx = 0
        t_idx = 0

        for elem in doc.element.body:
            if elem.tag.endswith('p'):
                if p_idx < len(doc.paragraphs):
                    paragraph = doc.paragraphs[p_idx]
                    p_idx += 1

                    text = paragraph.text.strip()
                    if not text:
                        continue

                    style_name = paragraph.style.name
                    if 'Heading' in style_name:
                        try:
                            level = int(style_name.split(' ')[-1])
                        except ValueError:
                            level = 1
                        md_lines.append(f"{'#' * level} {text}\n")
                    else:
                        md_lines.append(f"{text}\n")

            elif elem.tag.endswith('tbl'):
                if t_idx < len(doc.tables):
                    table = doc.tables[t_idx]
                    t_idx += 1
                    md_table = self._docx_table_to_md(table)
                    md_lines.append(md_table)
                    md_lines.append("\n")

        return "".join(md_lines)

    def _docx_table_to_md(self, table) -> str:
        """将 Word 表格转换为 Markdown 表格"""
        md_table = []
        if not table.rows:
            return ""

        headers = [cell.text.strip() for cell in table.rows[0].cells]
        md_table.append("| " + " | ".join(headers) + " |")
        md_table.append("| " + " | ".join(["---"] * len(headers)) + " |")

        for row in table.rows[1:]:
            row_data = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            md_table.append("| " + " | ".join(row_data) + " |")

        return "\n".join(md_table) + "\n"

    def _parse_pdf(self, file_path: Path) -> str:
        """解析文字版 PDF - 多进程并发优化版"""
        start_t = time.time()
        try:
            doc = pymupdf.open(file_path)
            num_pages = len(doc)
            doc.close()
        except Exception as e:
            logger.error(f"无法打开 PDF 文件 {file_path}: {e}")
            raise

        if num_pages == 0:
            return ""

        # 小文件（少于 10 页）直接单线程解析，减少进程创建开销
        if num_pages < 10:
            return self._parse_pdf_single(file_path)

        # 配置参数
        h2t_config = {
            "ignore_links": self.h2t.ignore_links,
            "ignore_images": self.h2t.ignore_images,
            "body_width": self.h2t.body_width,
            "ignore_emphasis": self.h2t.ignore_emphasis,
        }

        # 决定进程数：根据 PDF 页数动态调整，每 50 页增加一个进程
        # 最小 2 个进程，上限为 min(物理核数, 8)
        cpu_count = os.cpu_count() or 1
        dynamic_workers = (num_pages + 49) // 50
        max_workers = min(cpu_count, 8, max(2, dynamic_workers))
        
        # 任务分片
        pages_per_worker = (num_pages + max_workers - 1) // max_workers
        
        logger.info(f"开始多进程解析 PDF: {file_path.name}, 总页数: {num_pages}, 进程数: {max_workers}")
        
        results = [None] * max_workers
        futures = {}
        
        # 动态超时：基础 5 分钟 + 每 100 页增加 5 分钟，上限 30 分钟
        timeout_seconds = min(1800, 300 + (num_pages // 100) * 300)
        
        # 使用 ProcessPoolExecutor 进行并行解析
        # 显式管理 executor 以便在超时时能够非阻塞地关闭
        executor = ProcessPoolExecutor(max_workers=max_workers)
        try:
            for i in range(max_workers):
                start_page = i * pages_per_worker
                end_page = min((i + 1) * pages_per_worker, num_pages)
                if start_page < end_page:
                    future = executor.submit(
                        _parse_pdf_pages_worker, 
                        str(file_path), 
                        start_page, 
                        end_page, 
                        h2t_config
                    )
                    futures[future] = i

            # 收集结果
            for future in as_completed(futures, timeout=timeout_seconds):
                idx = futures[future]
                results[idx] = future.result()
                
        except TimeoutError:
            logger.error(f"PDF 解析超时 ({timeout_seconds}s): {file_path.name}")
            # 立即非阻塞关闭，尝试撤销尚未开始的任务
            executor.shutdown(wait=False, cancel_futures=True)
            # 在 Windows 上，我们需要主动结束这些可能卡死的子进程，但由于 ProcessPoolExecutor 
            # 封装较深，这里至少保证主线程不被 with 语句的隐式 shutdown 阻塞
            raise TimeoutError(f"PDF 解析超时 ({timeout_seconds}s)，文档可能过大或过于复杂")
        except Exception as e:
            logger.error(f"PDF 解析过程中发生错误: {str(e)}")
            executor.shutdown(wait=False)
            raise
        finally:
            # 正常情况下也需要关闭
            executor.shutdown(wait=False)

        duration = time.time() - start_t
        logger.info(f"PDF 多进程解析完成: {file_path.name}, 总耗时: {duration:.2f}s")
        
        # 过滤掉可能的 None 并合并
        return "\n\n".join([r for r in results if r is not None])

    def _parse_pdf_single(self, file_path: Path) -> str:
        """解析 PDF 逻辑，使用更稳健的 blocks 模式"""
        doc = pymupdf.open(file_path)
        md_content = []
        for page in doc:
            try:
                blocks = page.get_text("blocks")
                blocks.sort(key=lambda b: (b[1], b[0]))
                page_text = "\n\n".join([b[4].strip() for b in blocks if b[4].strip()])
                
                # 兜底
                if not page_text:
                    html_content = page.get_text("html")
                    page_text = self.h2t.handle(html_content)
                md_content.append(self._format_page_block(page.number + 1, page_text))
            except Exception as e:
                logger.warning(f"单线程解析 PDF 页面失败，尝试 HTML 模式: {e}")
                html_content = page.get_text("html")
                md_content.append(self._format_page_block(page.number + 1, self.h2t.handle(html_content)))
        doc.close()
        return "\n\n".join(md_content)

    def _parse_html(self, file_path: Path) -> str:
        """解析 .html 文件"""
        with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
            html_content = f.read()

        if not html_content.strip():
            return ""

        soup = BeautifulSoup(html_content, 'html.parser')

        for script in soup(["script", "style", "noscript", "header", "footer"]):
            script.decompose()

        block_tags = {'p', 'div', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6',
                      'li', 'ul', 'ol', 'tr', 'td', 'th', 'br', 'hr',
                      'article', 'section', 'aside', 'main', 'nav'}

        text_parts = []

        def extract_text_recursive(element):
            for child in element.children:
                if isinstance(child, str):
                    text = child.strip()
                    if text:
                        text_parts.append(text)
                else:
                    tag_name = child.name
                    extract_text_recursive(child)
                    if tag_name in block_tags:
                        text_parts.append("\n")

        root = soup.body if soup.body else soup
        extract_text_recursive(root)

        raw_text = "".join(text_parts)

        if not raw_text.strip():
            raw_text = soup.get_text(separator='\n')

        clean_text = re.sub(r'\n{3,}', '\n\n', raw_text)
        return clean_text

    def _parse_md(self, file_path: Path) -> str:
        """读取 .md 文件"""
        with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
            content = f.read()
        return content

    def _parse_txt(self, file_path: Path) -> str:
        """读取 .txt 文件"""
        with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
            return f.read()


# 全局文件解析器实例
file_parser = FileParser()
