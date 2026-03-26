import os
import sys
import logging
import traceback  # 用于获取完整的异常堆栈信息
import re  # 新增：用于正则清理多余的换行
from pathlib import Path
from typing import Optional, Callable
from multiprocessing import Pool, cpu_count
from tqdm import tqdm

# 格式化库
import fitz as pymupdf
import docx
import html2text
from bs4 import BeautifulSoup  # 新增：用于精细化的 HTML 解析

# 配置日志
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)


class FileParser:
    """单文件解析器：将不同格式文件转换为 Markdown 格式的文本"""

    def __init__(self):
        # 配置 HTML 转 Markdown 的参数 (主要用于 PDF 提取后的转换)
        self.h2t = html2text.HTML2Text()
        self.h2t.ignore_links = False
        self.h2t.ignore_images = True  # RAG 场景通常关注文本，忽略图片链接
        self.h2t.body_width = 0  # 不自动换行
        self.h2t.ignore_emphasis = False  # 保留加粗等强调

    def parse(self, file_path: Path) -> str:
        """根据后缀分发解析任务"""
        suffix = file_path.suffix.lower()
        try:
            if suffix == '.docx':
                return self._parse_docx(file_path)
            elif suffix == '.pdf':
                return self._parse_pdf(file_path)
            elif suffix == '.html':
                return self._parse_html(file_path)
            elif suffix == '.md':
                return self._parse_md(file_path)
            elif suffix == '.txt':  # 新增 .txt 支持
                return self._parse_txt(file_path)
            else:
                logger.warning(f"不支持的文件格式: {file_path}")
                return ""
        except Exception as e:
            error_detail = traceback.format_exc()
            logger.error(f"解析文件失败 {file_path}:\n{error_detail}")
            return ""

    def _parse_docx(self, file_path: Path) -> str:
        """解析 .docx 文件，保留标题层级和表格"""
        try:
            doc = docx.Document(file_path)
            md_lines = []

            # 使用计数器来解决底层 XML 索引与高级对象列表不匹配的问题
            p_idx = 0
            t_idx = 0

            for elem in doc.element.body:
                # 判断是否为段落
                if elem.tag.endswith('p'):
                    # 确保不越界
                    if p_idx < len(doc.paragraphs):
                        paragraph = doc.paragraphs[p_idx]
                        p_idx += 1

                        text = paragraph.text.strip()
                        if not text:
                            continue

                        # 简单的标题映射 (Word 样式 'Heading 1' -> '#')
                        style_name = paragraph.style.name
                        if 'Heading' in style_name:
                            try:
                                level = int(style_name.split(' ')[-1])  # "Heading 1" -> 1
                            except ValueError:
                                level = 1
                            md_lines.append(f"{'#' * level} {text}\n")
                        else:
                            md_lines.append(f"{text}\n")

                # 判断是否为表格
                elif elem.tag.endswith('tbl'):
                    if t_idx < len(doc.tables):
                        table = doc.tables[t_idx]
                        t_idx += 1
                        md_table = self._docx_table_to_md(table)
                        md_lines.append(md_table)
                        md_lines.append("\n")

            return "".join(md_lines)
        except Exception as e:
            error_detail = traceback.format_exc()
            logger.error(f"解析 DOCX 文件失败 {file_path}:\n{error_detail}")
            raise

    def _docx_table_to_md(self, table) -> str:
        """将 Word 表格转换为 Markdown 表格"""
        md_table = []
        if not table.rows:
            return ""

        # 表头
        headers = [cell.text.strip() for cell in table.rows[0].cells]
        md_table.append("| " + " | ".join(headers) + " |")
        md_table.append("| " + " | ".join(["---"] * len(headers)) + " |")

        # 表体
        for row in table.rows[1:]:
            row_data = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
            md_table.append("| " + " | ".join(row_data) + " |")

        return "\n".join(md_table) + "\n"

    def _parse_pdf(self, file_path: Path) -> str:
        """
        解析文字版 PDF
        使用 HTML 模式提取，再转 Markdown
        """
        try:
            if not file_path.exists():
                raise FileNotFoundError(f"PDF 文件不存在: {file_path}")

            if not os.access(file_path, os.R_OK):
                raise PermissionError(f"无法读取 PDF 文件: {file_path}")

            doc = pymupdf.open(file_path)
            md_content = []

            for page in doc:
                # 使用 "html" 模式，兼容性更好
                html_content = page.get_text("html")
                # 利用 html2text 转换
                md_text = self.h2t.handle(html_content)
                md_content.append(md_text)

            doc.close()

            # 合并并清理多余空行
            full_text = "\n\n".join(md_content)
            return full_text
        except Exception as e:
            error_detail = traceback.format_exc()
            logger.error(f"解析 PDF 文件失败 {file_path}:\n{error_detail}")
            raise

    def _parse_html(self, file_path: Path) -> str:
        """
        解析 .html 文件
        修改：按照 HTML 的换行标签（块级元素和 br）提取文字部分
        """
        try:
            # 修改点1：增加 errors='replace' 防止编码错误导致读取中断
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                html_content = f.read()

            if not html_content.strip():
                return ""

            # 使用 BeautifulSoup 解析
            soup = BeautifulSoup(html_content, 'html.parser')

            # 移除 script 和 style 标签及其内容
            for script in soup(["script", "style", "noscript", "header", "footer"]):
                script.decompose()

            # 定义被视为换行的块级标签
            block_tags = {'p', 'div', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6',
                          'li', 'ul', 'ol', 'tr', 'td', 'th', 'br', 'hr',
                          'article', 'section', 'aside', 'main', 'nav'}

            text_parts = []

            def extract_text_recursive(element):
                for child in element.children:
                    # 处理 NavigableString (文本内容)
                    if isinstance(child, str):
                        text = child.strip()
                        if text:
                            text_parts.append(text)
                    # 处理 Tag (HTML 标签)
                    else:
                        tag_name = child.name
                        # 递归处理子节点内容
                        extract_text_recursive(child)

                        # 如果是块级标签或 br 标签，处理完内容后添加换行
                        if tag_name in block_tags:
                            text_parts.append("\n")

            # 从 body 开始解析，如果没有 body 标签则解析整个文档
            root = soup.body if soup.body else soup
            extract_text_recursive(root)

            # 将提取的文本片段合并
            raw_text = "".join(text_parts)

            # 修改点2：如果自定义逻辑提取不到内容，回退到标准 get_text 方法
            if not raw_text.strip():
                # 降级策略：使用最通用的提取方式，尽量保证不为空
                raw_text = soup.get_text(separator='\n')
            
            # 清理：将连续超过2个的换行符替换为2个（Markdown 段落分隔）
            clean_text = re.sub(r'\n{3,}', '\n\n', raw_text)

            return clean_text

        except Exception as e:
            error_detail = traceback.format_exc()
            logger.error(f"解析 HTML 文件失败 {file_path}:\n{error_detail}")
            raise

    def _parse_md(self, file_path: Path) -> str:
        """读取 .md 文件并简单清洗（去除多余空行）"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                content = f.read()
            return content
        except Exception as e:
            error_detail = traceback.format_exc()
            logger.error(f"读取 MD 文件失败 {file_path}:\n{error_detail}")
            raise

    def _parse_txt(self, file_path: Path) -> str:
        """读取 .txt 文件"""
        try:
            # 使用 errors='replace' 防止部分特殊编码字符导致程序崩溃
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                return f.read()
        except Exception as e:
            error_detail = traceback.format_exc()
            logger.error(f"读取 TXT 文件失败 {file_path}:\n{error_detail}")
            raise


def _process_file_wrapper(args):
    """
    多进程 Worker 包装函数
    """
    input_path, output_dir = args
    parser = FileParser()

    try:
        # 解析内容
        md_content = parser.parse(input_path)

        if not md_content:
            logger.warning(f"文件内容为空: {input_path}")
            return None

        # 构建输出路径
        # 修改：按照需求转化为带格式的txt文件 (后缀改为 .txt)
        output_filename = f"{input_path.stem}.txt"
        output_path = Path(output_dir) / output_filename

        # 写入文件
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(md_content)

        return output_path
    except Exception as e:
        error_detail = traceback.format_exc()
        logger.error(f"处理文件异常 {input_path}:\n{error_detail}")
        return None


class BatchConverter:
    """批量转换器：使用多进程加速"""

    def __init__(self, input_dir: str, output_dir: str, num_processes: int = None):
        self.input_dir = Path(input_dir)
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

        # 默认使用 CPU 核心数，但至少保留 1 个核心给系统
        self.num_processes = num_processes if num_processes else max(1, cpu_count() - 1)
        logger.info(f"初始化转换器，使用进程数: {self.num_processes}")

    def get_files(self) -> list[Path]:
        """收集所有需要解析的文件"""
        # 修改：增加 .txt 支持
        valid_extensions = {'.docx', '.pdf', '.md', '.html', '.txt'}
        files = []
        for f in self.input_dir.rglob('*'):
            if f.is_file() and f.suffix.lower() in valid_extensions:
                files.append(f)
        return files

    def run(self):
        """执行多进程转换"""
        files = self.get_files()
        total_files = len(files)

        if total_files == 0:
            logger.warning("未找到任何支持的文件。")
            return

        logger.info(f"开始批量转换，共 {total_files} 个文件...")

        # 准备参数列表
        tasks = [(f, self.output_dir) for f in files]

        # 使用进程池
        success_count = 0
        failed_files = []

        with Pool(processes=self.num_processes) as pool:
            # 使用 tqdm 显示进度条
            results = list(tqdm(
                pool.imap(_process_file_wrapper, tasks),
                total=total_files,
                desc="Converting files"
            ))

            # 统计结果
            for i, res in enumerate(results):
                if res:
                    success_count += 1
                else:
                    failed_files.append(files[i])

        logger.info(f"转换完成！成功: {success_count}/{total_files}")
        if failed_files:
            logger.warning(f"失败的文件: {failed_files}")
        logger.info(f"输出目录: {self.output_dir.absolute()}")


if __name__ == "__main__":
    current_dir = os.path.dirname(os.path.abspath(__file__))
    os.chdir(current_dir)

    # 使用示例
    INPUT_DIR = "../data/raw_data"
    OUTPUT_DIR = "../data/cleaned_txt"  # 修改输出目录名以体现 txt 输出

    import argparse
    parser_cli = argparse.ArgumentParser()
    parser_cli.add_argument("--proc", type=int, default=4, help="进程数")
    parser_cli.add_argument("--input", type=str, default=INPUT_DIR, help="输入目录")
    parser_cli.add_argument("--output", type=str, default=OUTPUT_DIR, help="输出目录")
    args = parser_cli.parse_args()

    converter = BatchConverter(args.input, args.output, num_processes=args.proc)
    converter.run()
